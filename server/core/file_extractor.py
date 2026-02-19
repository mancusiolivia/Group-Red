"""
File extraction utility for parsing uploaded documents
Supports PDF, TXT, and DOCX files
"""
import os
from typing import Optional
from fastapi import UploadFile, HTTPException


def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text content from uploaded file
    Supports: PDF, TXT, DOCX
    """
    filename = file.filename or ""
    file_extension = os.path.splitext(filename)[1].lower()
    
    try:
        # Read file content
        file_content = file.file.read()
        file.file.seek(0)  # Reset file pointer
        
        if file_extension == '.txt':
            # Plain text file
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Try other encodings
                try:
                    text = file_content.decode('latin-1')
                except:
                    text = file_content.decode('utf-8', errors='ignore')
            return text.strip()
        
        elif file_extension == '.pdf':
            # PDF file
            try:
                import PyPDF2
                from io import BytesIO
                
                pdf_file = BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                text_parts = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as e:
                        print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                        continue
                
                if not text_parts:
                    raise HTTPException(
                        status_code=400,
                        detail="Could not extract text from PDF. The file may be image-based or corrupted."
                    )
                
                return "\n\n".join(text_parts).strip()
            except ImportError:
                raise HTTPException(
                    status_code=500,
                    detail="PDF extraction requires PyPDF2. Please install it: pip install PyPDF2"
                )
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Error extracting text from PDF: {str(e)}"
                )
        
        elif file_extension in ['.docx', '.doc']:
            # Word document
            try:
                from docx import Document
                from io import BytesIO
                
                doc_file = BytesIO(file_content)
                doc = Document(doc_file)
                
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                if not text_parts:
                    raise HTTPException(
                        status_code=400,
                        detail="Could not extract text from document. The file may be empty or corrupted."
                    )
                
                return "\n\n".join(text_parts).strip()
            except ImportError:
                raise HTTPException(
                    status_code=500,
                    detail="DOCX extraction requires python-docx. Please install it: pip install python-docx"
                )
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Error extracting text from document: {str(e)}"
                )
        
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_extension}. Supported types: .pdf, .txt, .docx, .doc"
            )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error processing file: {str(e)}"
        )


def summarize_text(text: str, max_length: int = 5000) -> str:
    """
    Summarize text if it's too long for LLM context
    For now, just truncate. Could be enhanced with actual summarization.
    """
    if len(text) <= max_length:
        return text
    
    # Truncate and add note
    truncated = text[:max_length]
    return f"{truncated}\n\n[Note: Document was truncated. Original length: {len(text)} characters]"


def extract_topics_from_content(content: str, num_topics: int = 1) -> list:
    """
    Extract distinct topics from uploaded content
    Uses simple heuristics: sections, headings, paragraphs
    Returns a list of topics (strings)
    """
    if not content or not content.strip():
        return []
    
    # Split content into potential topic sections
    # Look for headings (lines that are short and end with colon or are all caps)
    lines = content.split('\n')
    topics = []
    
    # Method 1: Look for headings (lines that are short, bold indicators, or section markers)
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        # Skip empty lines
        if not line_stripped:
            continue
        
        # Check if line looks like a heading:
        # - Short line (less than 100 chars)
        # - Ends with colon
        # - All caps (potential heading)
        # - Starts with numbers (like "1. Topic", "Chapter 1", etc.)
        is_heading = (
            len(line_stripped) < 100 and (
                line_stripped.endswith(':') or
                line_stripped.isupper() or
                (line_stripped[0].isdigit() and len(line_stripped.split()) < 10) or
                line_stripped.startswith('#') or
                line_stripped.startswith('Chapter') or
                line_stripped.startswith('Section')
            )
        )
        
        if is_heading:
            # Clean up the heading
            topic = line_stripped.replace(':', '').replace('#', '').strip()
            if topic and len(topic) > 3:  # Must be meaningful
                topics.append(topic)
    
    # Method 2: If we don't have enough headings, split by paragraphs and take first sentences
    if len(topics) < num_topics:
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if not para.strip():
                continue
            # Take first sentence or first 50 chars as topic
            first_sentence = para.split('.')[0].strip()
            if first_sentence and len(first_sentence) > 10 and len(first_sentence) < 100:
                if first_sentence not in topics:
                    topics.append(first_sentence)
                    if len(topics) >= num_topics:
                        break
    
    # Method 3: If still not enough, split content into chunks
    if len(topics) < num_topics:
        # Split content into roughly equal chunks
        chunk_size = len(content) // num_topics
        for i in range(num_topics):
            start_idx = i * chunk_size
            end_idx = start_idx + chunk_size if i < num_topics - 1 else len(content)
            chunk = content[start_idx:end_idx].strip()
            if chunk:
                # Take first meaningful line from chunk
                first_line = chunk.split('\n')[0].strip()
                if first_line and len(first_line) > 10 and len(first_line) < 150:
                    if first_line not in topics:
                        topics.append(first_line)
    
    # Return only the requested number of topics
    return topics[:num_topics] if topics else [content[:200]]  # Fallback to first 200 chars if no topics found
